import os
from PyQt5.QtCore import QThread, QObject, pyqtSignal
from bs4 import BeautifulSoup
from icecream import ic
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import docx
from common import *
import json
from requests.exceptions import *
from PIL import Image
from skimage import io
import shutil

# Global variable
DEBUG_VAR = True
# f_config = open("config.json", "r", encoding="utf-8")
# CONFIG_JSON = json.load(f_config)
# f_config.close()


def generateName(num: int, l: int) -> str:
    """Generate name from the number with length"""
    result_str = "0"*l + str(num)
    return result_str[-1*l:]


def generateChapterLink(chapter_str: str) -> str:
    """Generate chapter link from server mangasee123.com"""
    index = ""

    idexStr = int(chapter_str[0])

    if (idexStr != 1):
        index = '-index-' + idexStr

    chapter = int(chapter_str[1:-1])

    odd = ""

    odd_str = int(chapter_str[-1])

    if odd_str != 0:
        odd = "." + str(odd_str)

    return "-chapter-" + str(chapter) + odd + index


def generateChapterImg(chapter_str: str) -> str:
    """Generate chapter image file name server mangasee123.com"""
    chapter_str = str(chapter_str)
    chapter = chapter_str[1:-1]
    odd = chapter_str[-1]

    if odd == "0":
        return chapter
    else:
        return chapter + "." + odd


def generatePageImg(page: int) -> str:
    """Generate page image file name"""
    result_page = "000" + str(page)
    return result_page[-3:]


def downloadImage(link: str, server: str, file: str, count: int):
    """Download image from link"""

    if os.path.exists(file):

        # Check file exist and is a valid image
        # if image valid return 200
        # else remove file and download again until count = 3
        try:
            img = Image.open(file)
            img.verify()
            img = io.imread(file)
            return 200
        except:
            os.remove(file)
            if count < 3:
                return downloadImage(link, server, file,  count+1)

    else:
        if count < 3:
            try:
                r = requests.get(link.replace("\n", ""), headers={
                    'User-agent': 'Mozilla/5.0', 'Referer': server}, timeout=(3, 5))

                flag = False  # Check download success or not
                down_code = 200
                with open(file, "wb") as fd:
                    down_code = r.status_code
                    if (down_code != 200):
                        flag = True
                    else:
                        fd.write(r.content)

                if flag:
                    # If download fail and status code is not 404
                    # true: remove file and download again until count = 3
                    # false: return 404
                    if down_code == 404:
                        os.remove(file)
                        return 404
                    return downloadImage(link, server, file, count+1)
                else:
                    return downloadImage(link, server, file, count)
            except Exception as e:
                DEBUG_VAR and print("Error: ", e)
                return 400


class RenameFolder(QObject):
    finished = pyqtSignal()
    progress = pyqtSignal(int)

    def __init__(self, link):
        QObject.__init__(self)
        self.link = link

    def run(self):

        chapter = []
        f = open("resource/list_files.txt", "r", encoding="utf-8")

        for x in f:
            chapter.append(x)

        chapter_new = []
        f = open("resource/list_files_new.txt", "r", encoding="utf-8")

        for x in f:
            chapter_new.append(x)

        self.progress.emit(len(chapter_new))

        for old, new in zip(chapter, chapter_new):
            os.rename(self.link+"/"+old.replace("\n", ""),
                      self.link+"/"+new.replace("\n", ""))
            self.progress.emit(1)

        self.finished.emit()


def writeChapter(chaps, title, content, targetFile=None, file_type='docx'):
    if file_type == 'docx' and targetFile and isinstance(targetFile, docx.document.Document):
        p = targetFile.add_heading(title.text.strip(), level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = targetFile.add_paragraph(break_chapter_str)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = targetFile.add_paragraph(content.strip())
        targetFile.add_page_break()
        return targetFile
    elif file_type == 'txt':
        f = open("resource/"+str(chaps)+'.txt', encoding='utf8', mode='w+')
        f.write(title.text.strip()+"\n")
        f.write(content.strip()+'\n')
        f.close()


class DownloadNovel(QObject):
    finished = pyqtSignal(tuple)
    progress = pyqtSignal(tuple)

    def __init__(self, link, start, end, novelName, author, file_name, server, file_type):
        QObject.__init__(self)
        self.link, self.start, self.end, self.novelName, self.author, self.file_name, self.server, self.file_type = link, start, end, novelName, author, file_name, server, file_type

    def run(self):

        if self.file_name == "":
            filename = "resource/"+self.novelName + ' chap' + \
                str(self.start)+'_'+str(self.end)
        else:
            filename = "resource/"+self.file_name.strip()

        if self.file_type == 0:
            document = Document()

        print(self.server)

        try:
            if self.server == servers_novel["metruyencv"]:
                count = 0
                for i in range(self.start, self.end+1):
                    count += 1
                    r = requests.get(self.link+str(i))

                    soup = BeautifulSoup(r.content, 'html.parser')

                    title = soup.find(
                        "h2", class_=re.compile("text-balance"))

                    ic(title)
                    ic(title.text.strip())

                    content = soup.find_all(_class="break-words")[0]

                    content_text = content.decode_contents()
                    content_text = content_text.replace("<br/>", "\n")

                    soup = BeautifulSoup(content_text, 'html.parser')
                    content_text = soup.text

                    if "— QUẢNG CÁO —" in content_text:
                        content_text = content_text.replace(
                            "— QUẢNG CÁO —", "")
                    if self.file_type == 0:
                        document = writeChapter(
                            i, title, content_text, document)
                    else:
                        writeChapter(i, title, content_text, file_type='txt')

                    self.progress.emit((title.text.strip(), int(
                        count*100/(self.end-self.start+1))))

                self.finished.emit((filename+'.docx', 200))
            elif self.server == servers_novel["sstruyen"]:

                count = 0
                for i in range(self.start, self.end+1):
                    count += 1
                    r = requests.get(self.link+str(i))

                    soup = BeautifulSoup(r.content, 'html.parser')

                    title = soup.find_all(
                        class_=re.compile("rv-chapt-title"))[0]
                    ic(title.text.strip())

                    content = soup.find_all(class_=re.compile("container1"))

                    content = str(content[0]).replace("<br/>", "\n")
                    soup = BeautifulSoup(content, 'html.parser')
                    content_text = soup.text.strip()

                    p = document.add_heading(title.text.strip(), level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = document.add_paragraph(break_chapter_str)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = document.add_paragraph(content_text.strip())
                    document.add_page_break()

                    self.progress.emit((title.text.strip(), int(
                        count*100/(self.end-self.start+1))))
                self.finished.emit((filename+'.docx', 200))

            elif self.server == servers_novel["trumtruyen"]:

                count = 0
                for i in range(self.start, self.end+1):
                    count += 1
                    r = requests.get(self.link+str(i))

                    QThread.sleep(1)

                    while r.status_code != 200:
                        r = requests.get(self.link+str(i))
                        QThread.sleep(1)

                    soup = BeautifulSoup(r.content, 'html.parser')

                    title = soup.find_all(
                        class_=re.compile("chapter-title"))[0]

                    content = soup.find_all(id="chapter-c")

                    content = str(content[0]).replace("<br/>", "\n")
                    soup = BeautifulSoup(content, 'html.parser')
                    content_text = soup.text.strip()

                    p = document.add_heading(title.text.strip(), level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = document.add_paragraph(break_chapter_str)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = document.add_paragraph(content_text.strip())
                    document.add_page_break()

                    self.progress.emit((title.text.strip(), int(
                        count*100/(self.end-self.start+1))))

                self.finished.emit((filename+'.docx', 200))

            elif self.server == servers_novel["truyenfull"]:
                count = 0
                for i in range(self.start, self.end+1):
                    count += 1
                    r = requests.get(self.link+str(i))

                    soup = BeautifulSoup(r.content, 'html.parser')

                    title = soup.find_all(
                        class_=re.compile("chapter-title"))[0]

                    content = soup.find_all(id="chapter-c")

                    content = str(content[0]).replace("<br/>", "\n")
                    soup = BeautifulSoup(content, 'html.parser')
                    content_text = soup.text.strip()

                    p = document.add_heading(title.text.strip(), level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = document.add_paragraph(break_chapter_str)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = document.add_paragraph(content_text.strip())
                    document.add_page_break()

                    self.progress.emit((title.text.strip(), int(
                        count*100/(self.end-self.start+1))))

                self.finished.emit((filename+'.docx', 200))

        except HTTPError:
            self.finished.emit(
                ('No internet connection!!! Please connect internet!', 400))
            if self.file_type == 0:
                core_properties = document.core_properties
                core_properties.author = self.author
                core_properties.comments = "Generated by Crawl Manga - An Đào"
                document.save(filename+".docx")

        except ConnectionError:
            self.finished.emit(
                ('No internet connection!!! Please connect internet!', 400))
            if self.file_type == 0:
                core_properties = document.core_properties
                core_properties.author = self.author
                core_properties.comments = "Generated by Crawl Manga - An Đào"
                document.save(filename+".docx")

        except Exception as e:
            self.finished.emit(('Error!!!', -1))
            DEBUG_VAR and print("error: ", e)
            if self.file_type == 0:
                core_properties = document.core_properties
                core_properties.author = self.author
                core_properties.comments = "Generated by Crawl Manga - An Đào"
                document.save(filename+".docx")

        if self.file_type == 0:
            core_properties = document.core_properties
            core_properties.author = self.author
            core_properties.comments = "Generated by Crawl Manga - An Đào"
            document.save(filename+".docx")


class GetChapterLink(QObject):
    finished = pyqtSignal(tuple)
    progress = pyqtSignal(dict)

    def __init__(self, link, server):
        QObject.__init__(self)
        self.link, self.server = link, server

    def run(self):

        r = requests.get(self.link.strip(), headers={
            'User-agent': 'Mozilla/5.0'})

        links = []

        self.progress.emit({"title": "Stage 1", "percent": -1})

        link_s = self.link.split("/")

        ic(link_s)

        l_server = "/".join(link_s[0:3])
        ic(l_server)

        if os.path.exists('resource/chapter.json'):
            f_json = open('resource/chapter.json', 'r', encoding='utf8')
            chapters_obj = json.load(f_json)
            f_json.close()
        else:
            chapters_obj = {}

            chapters_obj["server"] = l_server

            chapters_obj["chapters"] = {}
            chapters_obj["links"] = []

        try:
            if self.server == "nettruyen":

                mang_name_link = link_s[-1]
                mang_name_link = mang_name_link.split('-')
                mang_name_link.pop()
                mang_name_link = "-".join(mang_name_link)

                self.keyword = mang_name_link

                htmlSource = r.content
                soup = BeautifulSoup(htmlSource, 'html.parser')

                link_chapters = soup.find_all(id="nt_listchapter")
                link_chapters = link_chapters[0].find_all(
                    href=re.compile(self.keyword))

                ic(link_chapters)

                for idx, link in enumerate(link_chapters):
                    if link['href'] != "#":
                        links.insert(0, link['href'])
                        self.progress.emit(
                            {"title": link['href'], "percent": int((idx+1)*100/len(link_chapters))})

                self.progress.emit({"title": "Stage 2", "percent": -1})

                count = 0
                for link in links:

                    count += 1

                    split_link_chap = link.split('/')
                    id_chap_link = split_link_chap[-2]+'/'+split_link_chap[-1]

                    if id_chap_link in chapters_obj["links"]:
                        ic(id_chap_link)
                        self.progress.emit(
                            {"title": "Skip down", "percent": int((count)*100/len(links))})
                    else:

                        r = requests.get(link.strip(), headers={
                            'User-agent': 'Mozilla/5.0'})

                        htmlSource = r.content

                        soup = BeautifulSoup(htmlSource, 'html.parser')

                        title = soup.find('title')
                        title = title.text.split(" Next Chap ")[0].strip()

                        chap = title.split(" ")[-1]
                        odd = -1
                        if len(chap.split(".")) >= 2:
                            odd = chap.split(".")[-1]
                            chap = chap.split(".")[0]

                        title = "Chapter " + generateName(chap, 4)
                        if odd != -1:
                            title = title + "." + odd

                        ic(title)

                        chapters_obj["chapters"][title] = []

                        divs = soup.find_all("div", id=re.compile('page_'))

                        if len(divs) == 0:
                            divs = soup.find_all(
                                "div", class_=re.compile('page-'))

                        self.progress.emit(
                            {"title": title, "percent": int((count)*100/len(links))})

                        # target field get img url
                        target_field = "data-src"

                        for idxx, div in enumerate(divs):
                            img = div.find('img')
                            if img[target_field] is None or img[target_field] == "":
                                raise Exception("Field wrong !!!")
                            else:
                                if "https:" not in img[target_field]:
                                    chapters_obj["chapters"][title].append(
                                        'https:'+img[target_field])
                                else:
                                    chapters_obj["chapters"][title].append(
                                        img[target_field])

                        if len(chapters_obj["chapters"][title]) > 0:
                            chapters_obj['links'].append(id_chap_link)

                with open("resource/chapter.json", "w+") as outfile:
                    outfile.write(json.dumps(chapters_obj, indent=2))

            else:
                r = requests.get(self.link.strip())

                f = open('resource/test.html', mode='w+', encoding='utf-8')
                f.write(r.text)
                f.close()

                chapters = []
                index_name = ""

                f = open('resource/test.html', mode='r+', encoding='utf-8')

                self.progress.emit({"title": "Stage 1", "percent": -1})
                for line in f.readlines():
                    if "vm.CurPathName = " in line:
                        cur_path_name = line.replace("vm.CurPathName = ", "")

                    if "vm.IndexName = " in line:
                        index_name = line.replace("vm.IndexName = ", "")
                        index_name = index_name.strip()
                        index_name = index_name.replace(
                            '"', '').replace(";", "")

                    if "vm.CHAPTERS =" in line:
                        chapters = json.loads(line.replace(
                            'vm.CHAPTERS = ', "").replace(";", ""))

                self.progress.emit({"title": "Stage 2", "percent": -1})

                if len(chapters) != 0:
                    for chap_idx, chap in enumerate(chapters):

                        id_chap_link = index_name + \
                            generateChapterLink(chap["Chapter"])

                        if id_chap_link in chapters_obj['links']:
                            ic(id_chap_link)
                            self.progress.emit(
                                {"title": "Skip down", "percent": int((chap_idx+1)/len(chapters)*100)})
                        else:

                            link = "https://mangasee123.com/read-online/" + id_chap_link + ".html"

                            r = requests.get(link)
                            f_tmp = open('tmp.html', 'w+', encoding='utf-8')
                            f_tmp.write(r.text)
                            f_tmp.close()
                            f_tmp = open('tmp.html', 'r+', encoding='utf-8')

                            cur_path_name = ""
                            for line in f_tmp.readlines():
                                if "vm.CurPathName = " in line:
                                    cur_path_name = line.replace(
                                        "vm.CurPathName = ", "").strip().replace(";", "").replace('"', "")
                                    break

                            chap_name = "Chapter " + \
                                generateChapterImg(chap["Chapter"])
                            chapters_obj["chapters"][chap_name] = []

                            self.progress.emit(
                                {"title": chap_name, "percent": int((chap_idx+1)/len(chapters)*100)})

                            for p_idx in range(1, int(chap["Page"])+1):
                                img_link = "https://{curPathName}/manga/{index_name}/{directory}{img}.png"
                                img_link = img_link.replace(
                                    "{curPathName}", cur_path_name)
                                img_link = img_link.replace(
                                    "{index_name}", index_name)
                                if chap["Directory"] != "":
                                    img_link = img_link.replace(
                                        "{directory}", chap["Directory"])
                                else:
                                    img_link = img_link.replace(
                                        "{directory}", "")
                                img_link = img_link.replace("{img}", generateChapterImg(
                                    chap["Chapter"])+"-"+generatePageImg(p_idx))

                                chapters_obj["chapters"][chap_name].append(
                                    img_link)

                            if len(chapters_obj["chapters"][chap_name]) > 0:
                                chapters_obj['links'].append(id_chap_link)

                    with open("resource/chapter.json", "w+") as outfile:
                        outfile.write(json.dumps(chapters_obj, indent=2))
        except Exception as e:
            ic(e)
            with open("resource/chapter.json", "w+") as outfile:
                outfile.write(json.dumps(chapters_obj, indent=2))
            self.finished.emit(("Error!", 400))

        self.finished.emit(("Success!!!", 200))


class DownloadImage(QObject):
    finished = pyqtSignal(tuple)
    progress = pyqtSignal(tuple)

    def __init__(self):
        QObject.__init__(self)

    def run(self):

        chap_f = open('resource/chapter.json', 'r', encoding='utf-8')

        chapter_obj = json.load(chap_f)

        keys_chap = list(chapter_obj["chapters"])

        crr_chap = keys_chap[0]

        crr_idx = 0

        if "prev_down" in chapter_obj.keys():
            crr_chap = chapter_obj["prev_down"]
            crr_idx = keys_chap.index(crr_chap)

        breakAll = False
        while crr_idx != len(keys_chap):

            crr_chap = keys_chap[crr_idx]

            self.progress.emit(
                (crr_chap, int((crr_idx+1)*100/(len(keys_chap)))))

            if not os.path.exists(crr_chap) and len(chapter_obj["chapters"][crr_chap]) != 0:
                os.mkdir(crr_chap)

            for idx, link in enumerate(chapter_obj["chapters"][crr_chap]):
                res = downloadImage(
                    link, chapter_obj['server'], crr_chap+"/"+generateName(idx+1, 3)+'.jpg', 0)

                if res == 200:
                    self.progress.emit(
                        (crr_chap + ' - ' + generateName(idx+1, 3)+'.jpg', int((crr_idx+1)*100/(len(keys_chap)))))
                elif res == 404:
                    continue
                else:
                    self.finished.emit(('Error', 400))
                    chapter_obj["prev_down"] = crr_chap
                    with open("resource/chapter.json", "w+") as outfile:
                        outfile.write(json.dumps(chapter_obj, indent=2))
                    return
            if breakAll:
                break
            crr_idx += 1

        self.finished.emit(('Success', 200))


class DownloadInfoComic(QObject):
    finished = pyqtSignal(tuple)
    progress = pyqtSignal(tuple)

    def __init__(self, link, isCover, isVolChapters):
        QObject.__init__(self)

        self.link = link
        self.isCover = isCover
        self.isVolChapters = isVolChapters

    def run(self):

        part_link = self.link.split("/")

        if len(part_link) > 1:
            manga_idx = part_link[-2]

        self.progress.emit(('Stage 1: Download comic info', 0))

        link_cminfo_api = "https://api.mangadex.org/manga/{manga_idx}?includes[]=artist&includes[]=author&includes[]=cover_art"
        link_cminfo_api = link_cminfo_api.replace("{manga_idx}", manga_idx)

        r = requests.get(link_cminfo_api, headers={
            'User-agent': 'Mozilla/5.0'})

        data = r.json()["data"]

        title = data["attributes"]['title']
        comic_name = title[list(title)[0]]
        ic(comic_name)

        author = ""

        author_data = data["relationships"]
        for rel in author_data:
            if rel['type'] == 'author':
                author = rel['attributes']['name']

        self.progress.emit(
            ('Complete stage 1: '+str(comic_name)+' - '+author, 100))

        results = []

        results.append([comic_name, author])

        if self.isVolChapters:
            self.progress.emit(('Stage 2: Download volume info', 0))

            offset = 0

            link_volInfo_api = "https://api.mangadex.org/manga/{manga_idx}/feed?includes[]=scanlation_group&includes[]=user&order[volume]=asc&order[chapter]=asc&offset={offset}"

            link_volInfo_api = link_volInfo_api.replace(
                "{manga_idx}", manga_idx)

            total = 0

            vols = dict()

            count_repeat = 0

            while True:

                tmp = link_volInfo_api.replace("{offset}", str(offset))

                try:
                    r = requests.get(tmp,  headers={
                        'User-agent': 'Mozilla/5.0', 'Referer': 'https://mangadex.org/'}, timeout=(3, 5))
                    ic(tmp)
                    data = r.json()
                    count_repeat = 0

                    total = data['total']
                    for chap in data["data"]:
                        crr_vol = chap["attributes"]["volume"]
                        if chap["attributes"]["chapter"] is not None:
                            crr_chap = int(
                                float(chap["attributes"]["chapter"]))
                        else:
                            continue

                        if crr_vol not in vols.keys():
                            vols[crr_vol] = []
                        if crr_chap not in vols[crr_vol]:
                            vols[crr_vol].append(crr_chap)

                    offset += 100
                    if offset > total:
                        break
                except json.decoder.JSONDecodeError as errJson:
                    ic(errJson)
                    count_repeat = count_repeat+1
                    if count_repeat == 5:
                        self.finished.emit(("Something error happen!!!", -1))
                        return
                    else:
                        continue

            for idx, key in enumerate(vols.keys()):
                if key is not None:
                    self.progress.emit(('Volume '+str(key), int(
                        (idx+1)*100/(len(vols.keys())))))
                    if len(vols[key]) >= 2:
                        results.append([int(key), vols[key][0], vols[key][-1]])

        if self.isCover:

            self.progress.emit(('Stage 2: Download cover', 0))

            offset = 0
            link_cov_api = "https://api.mangadex.org/cover?order[volume]=asc&manga[]={manga_idx}&limit=100&offset={offset}"

            link_cov_api = link_cov_api.replace("{manga_idx}", manga_idx)

            if not os.path.exists('resource/covers'):
                os.mkdir("resource/covers")

            covers = []
            while True:
                crr_link = link_cov_api.replace("{offset}", str(offset))
                r = requests.get(crr_link)
                result = r.json()
                data = result["data"]
                for idx, cover in enumerate(data):
                    cover_link = "https://mangadex.org/covers/"
                    cover_link = cover_link + \
                        cover["relationships"][0]['id']+'/' + \
                        cover["attributes"]["fileName"]

                    tmp_vol = cover["attributes"]["volume"]

                    size_img = 3

                    if tmp_vol and "." in tmp_vol:
                        size_img = 4

                    filename = generateName(tmp_vol, size_img) + "." + \
                        cover["attributes"]["fileName"].split(".")[-1]

                    ic(cover_link)

                    covers.append({'link': cover_link, 'filename': filename})

                offset += 100
                if result['total'] < offset:
                    break

            for idx, cover in enumerate(covers):
                r = requests.get(cover['link'], headers={
                    'User-agent': 'Mozilla/5.0', 'Referer': "https://mangadex.org/"}, timeout=(3, 5))

                if r.status_code != 200:
                    self.finished.emit(("Something error happen!!!", -1))
                    break
                else:
                    with open("resource/covers"+"/"+cover['filename'], "wb") as fd:
                        if (r.status_code != 200):
                            ic("error")
                        else:
                            fd.write(r.content)
                            self.progress.emit((cover['link'], int(
                                (idx+1)*100/(len(covers)))))

        self.finished.emit((results, 200))


def createNameFolder(link, comic, vol, type_c):
    if type_c == 'folder':
        return link + '/' + comic + ' - Vol' + str(vol)
    elif type_c == 'info':
        return link + '/' + comic + ' - Vol' + str(vol) + '/ComicInfo.xml'
    elif type_c == 'progress':
        return comic + ' - Vol' + str(vol)
    elif type_c == 'fol_cov':
        return (link + '/' + comic + ' - Vol' + str(vol) + '/00_Cover', link + '/' + comic + ' - Vol' + str(vol) + '/00_Cover/0000.jpg')
    else:
        return link + '/' + comic + ' - Vol' + str(vol)+'/' + '0000.jpg'


class ArchiveComic(QObject):
    finished = pyqtSignal(tuple)
    progress = pyqtSignal(tuple)

    def __init__(self, src, des, vol_chap, cover="", isMakeFolder=False):
        QObject.__init__(self)

        self.src, self.des, self.vol_chap, self.cover, self.isMakeFolder = src, des, vol_chap, cover, isMakeFolder

    def run(self):

        xmlContaint = "<?xml version=\"1.0\" encoding=\"utf-8\"?>" +\
            "<ComicInfo xmlns:xsd=\"http://www.w3.org/2001/XMLSchema\" xmlns:xsi=\"http://www.w3.org/2001/XMLSchema-instance\">" +\
            "{content}" +\
            "</ComicInfo>"

        if self.src == "" or self.des == "" or self.vol_chap == "":
            self.finished.emit(('Please fill the information'))

        f = open(self.vol_chap, 'r', encoding='utf8')
        vol_chaps = json.load(f)
        f.close()

        self.progress.emit(('Stage 1: Create cover', 0))

        tmp_fol_name = vol_chaps['comic']
        if 'comic_s' in vol_chaps.keys():
            tmp_fol_name = vol_chaps['comic_s']

        count_vol = 0
        total_vol = vol_chaps['end_vol'] - vol_chaps['start_vol'] + 1
        for i in range(vol_chaps['start_vol'], vol_chaps['end_vol'] + 1):
            count_vol = count_vol + 1

            self.progress.emit((createNameFolder(self.des, tmp_fol_name, i, 'progress'), int(
                count_vol*100/(total_vol))))
            crr_fol = createNameFolder(
                self.des, tmp_fol_name, i, 'folder')
            if not os.path.exists(crr_fol):
                os.mkdir(crr_fol)

        self.progress.emit(('Stage 2: Comic info', 0))

        count_vol = 0
        for i in range(vol_chaps['start_vol'], vol_chaps['end_vol'] + 1):
            count_vol = count_vol + 1
            series = "<Series>"+vol_chaps['comic']+"</Series>\n"
            volume = "<Volume>"+str(i)+"</Volume>\n"
            writer = "<Writer>"+vol_chaps['author']+"</Writer>\n"
            final = xmlContaint
            final = final.replace("{content}", series+volume+writer)
            f_comic_info = open(createNameFolder(
                self.des, tmp_fol_name, i, 'info'), 'w+', encoding='utf8')
            f_comic_info.write(final)
            f_comic_info.close()
            self.progress.emit((createNameFolder(self.des, tmp_fol_name, i, 'progress'), int(
                count_vol*100/(total_vol))))

        self.progress.emit(('Stage 3: Move chapters - volume', 0))

        chapter_link_fol = os.listdir(self.src)

        count_vol = 0
        for i in range(vol_chaps['start_vol']-1, vol_chaps['end_vol']):
            count_vol = count_vol+1
            tmp_link_fol = createNameFolder(
                self.des, tmp_fol_name, i+1, 'folder')
            chapter_link_vol = []
            for chap_idx in range(vol_chaps['vols'][i]['start'], vol_chaps['vols'][i]['end']+1):
                tmp = [x for x in chapter_link_fol if "Chapter " +
                       generateName(str(chap_idx), 4) in x]
                chapter_link_vol.extend(tmp)
            self.progress.emit((createNameFolder(self.des, tmp_fol_name, i+1, 'progress'), int(
                count_vol*100/(total_vol))))
            for chap in chapter_link_vol:
                shutil.copytree(self.src+'/' + chap, tmp_link_fol + '/' + chap)

        if self.cover != "":
            self.progress.emit(('Move cover', 0))

            list_cov_files = os.listdir(self.cover)
            count_vol = 0
            for i in range(vol_chaps['start_vol'], vol_chaps['end_vol']+1):
                count_vol = count_vol+1
                target_cov = [
                    x for x in list_cov_files if generateName(i, 3) in x][0]
                ic(target_cov)

                tmp_link_fol, img = createNameFolder(
                    self.des, tmp_fol_name, i, 'fol_cov')

                if self.isMakeFolder:
                    if not os.path.exists(tmp_link_fol):
                        os.mkdir(tmp_link_fol)
                    shutil.copyfile(self.cover+'/' + target_cov, img)
                else:
                    img = createNameFolder(self.des, tmp_fol_name, i, "")
                    shutil.copyfile(self.cover+'/' + target_cov, img)

                self.progress.emit((createNameFolder(self.des, tmp_fol_name, i, 'progress'), int(
                    count_vol*100/(total_vol))))

        count_vol = 0
        self.progress.emit(('Stage 4: Archive comic', 0))
        for i in range(vol_chaps['start_vol'], vol_chaps['end_vol']+1):
            count_vol = count_vol + 1
            tmp_link_fol = createNameFolder(
                self.des, tmp_fol_name, i, 'folder')
            shutil.make_archive(tmp_link_fol, "zip", root_dir=tmp_link_fol)
            self.progress.emit((createNameFolder(self.des, tmp_fol_name, i, 'progress'), int(
                count_vol*100/(total_vol))))

        self.finished.emit(('Success', 200))
