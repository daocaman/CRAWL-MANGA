import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from bs4 import BeautifulSoup
from icecream import ic
import requests
from docx import Document
from win10toast import ToastNotifier
n = ToastNotifier()

document = Document()

start = 1791
end = 1792

vol = -1

link_novel = 'https://sstruyen.vn/mao-son-troc-quy-nhan/chuong-'

book_name = "Mao Sơn Tróc Quỷ Nhân"

author = "Thiên Hạ Bá Xướng"

page_break = "※----*-------⁛-------*----※"


# mode 1 : metruyencv
# mode 2: sstruyen
# mode 3: trumtruyen
# mode 4: truyenfull
mode = 2

try:
    if mode == 1:
        for i in range(start, end+1):
            r = requests.get(link_novel+str(i))

            soup = BeautifulSoup(r.content, 'html.parser')

            title = soup.find(class_="nh-read__title")

            ic(title.text.strip())

            content = soup.find_all(id="article")[0]

            content_text = content.decode_contents()
            content_text = content_text.replace("<br/>", "\n")

            soup = BeautifulSoup(content_text, 'html.parser')
            content_text = soup.text

            if "— QUẢNG CÁO —" in content_text:
                content_text = content_text.replace(
                    "— QUẢNG CÁO —", "")

            p = document.add_heading(title.text.strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(page_break)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(content_text.strip())
            document.add_page_break()
    elif mode == 2:

        for i in range(start, end+1):
            r = requests.get(link_novel+str(i))

            soup = BeautifulSoup(r.content, 'html.parser')

            title = soup.find_all(class_=re.compile("rv-chapt-title"))[0]
            ic(title.text.strip())

            content = soup.find_all(class_=re.compile("container1"))

            content = str(content[0]).replace("<br/>", "\n")
            soup = BeautifulSoup(content, 'html.parser')
            content_text = soup.text.strip()

            content_text = content_text.replace(u'\ufeff', '')
            content_text = content_text.replace(u'\x0b', '')

            p = document.add_heading(title.text.strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(page_break)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(content_text.strip())
            document.add_page_break()
    elif mode == 3:

        for i in range(start, end+1):
            r = requests.get(link_novel+str(i))

            time.sleep(0.5)
            # ic(r.status_code)
            while r.status_code != 200:
                r = requests.get(link_novel+str(i))
                time.sleep(0.5)

            soup = BeautifulSoup(r.content, 'html.parser')

            title = soup.find_all(class_=re.compile("chapter-title"))[0]
            ic(title.text.strip())

            content = soup.find_all(id="chapter-c")

            content = str(content[0]).replace("<br/>", "\n")
            soup = BeautifulSoup(content, 'html.parser')
            content_text = soup.text.strip()

            p = document.add_heading(title.text.strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(page_break)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(content_text.strip())
            document.add_page_break()
    else:
        for i in range(start, end+1):
            # r = requests.get(link_novel.replace('{0}', str(i)))
            r = requests.get(link_novel+str(i))

            soup = BeautifulSoup(r.content, 'html.parser')

            title = soup.find_all(class_=re.compile("chapter-title"))[0]
            ic(title.text.strip())

            content = soup.find_all(id="chapter-c")

            content = str(content[0]).replace("<br/>", "\n")
            soup = BeautifulSoup(content, 'html.parser')
            content_text = soup.text.strip()

            p = document.add_heading(title.text.strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(page_break)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p = document.add_paragraph(content_text.strip())
            document.add_page_break()

except Exception as e:
    ic(e)
    core_properties = document.core_properties
    core_properties.author = author
    core_properties.comments = "Generated by Crawl Manga - An Đào"
    if vol == -1:
        document.save(book_name + ' chap ' +
                    str(start)+'_'+str(end)+'.docx')
    else:
        document.save(book_name + '_tập '+str(vol)+'.docx')



core_properties = document.core_properties
core_properties.author = author
core_properties.comments = "Generated by Crawl Manga - An Đào"

if vol == -1:   
    document.save(book_name +' chap '+str(start)+'_'+str(end)+'.docx')
else:
    document.save(book_name + '_tập '+str(vol)+'.docx')
n.show_toast("Complete", "Getting novel complete!!!", duration=2)
